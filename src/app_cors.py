"""
Versión de app.py con CORS habilitado para funcionar con GitHub Pages y otros clientes.
Asegúrate de instalar: pip install flask-cors
"""
import os
import time
import uuid
import threading
import shutil
from docx import Document
from flask import Flask, request, redirect, url_for, render_template, send_from_directory, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

from splitter import split_audio
from transcriptor import Transcriptor
from diarizacion import Diarizador
from formato import Formateador
from ai_roles import AsignadorRoles


BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
UPLOAD_DIR = os.path.join(BASE_DIR, "audios", "uploads")
OUTPUT_DIR = os.path.join(BASE_DIR, "salidas")
DOCS_DIR = os.path.join(BASE_DIR, "docs")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

app = Flask(__name__, static_folder=DOCS_DIR, static_url_path="/static")
app.config["UPLOAD_FOLDER"] = UPLOAD_DIR

# Habilitar CORS para GitHub Pages
CORS(app, resources={
    r"/upload": {"origins": "*"},
    r"/job-status/*": {"origins": "*"},
    r"/salidas/*": {"origins": "*"}
})

JOBS = {}
JOBS_LOCK = threading.Lock()
TEMPLATE_NAME = 'index.html'


@app.route("/", methods=["GET"]) 
def index():
    """Sirve la interfaz web desde docs/index.html"""
    return send_file(os.path.join(DOCS_DIR, "index.html"))


@app.route("/style.css")
def style():
    """Sirve el CSS"""
    return send_file(os.path.join(DOCS_DIR, "style.css"), mimetype="text/css")


@app.route("/app.js")
def app_js():
    """Sirve el JavaScript"""
    return send_file(os.path.join(DOCS_DIR, "app.js"), mimetype="application/javascript")


@app.route("/upload", methods=["POST", "OPTIONS"])
def upload():
    if request.method == "OPTIONS":
        return "", 204
    
    files = request.files.getlist("files")
    if not files or all(f.filename == "" for f in files):
        return jsonify({"error": "No se envió ningún archivo"}), 400

    saved_files = []
    for f in files:
        if f.filename == "":
            continue
        filename = secure_filename(f.filename)
        saved_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        f.save(saved_path)
        saved_files.append(saved_path)

    if not saved_files:
        return jsonify({"error": "No se pudo guardar ningún archivo"}), 400

    job_id = str(uuid.uuid4())

    with JOBS_LOCK:
        JOBS[job_id] = {
            "status": "queued",
            "progress": 0,
            "total": 0,
            "result": None,
            "error": None,
            "txt_name": None,
            "docx_names": [],
            "roles": {},
        }

    def worker():
        try:
            with JOBS_LOCK:
                JOBS[job_id]["status"] = "splitting"

            all_results = []
            speaker_set = []
            total_chunks = 0
            processed_chunks = 0
            docx_names = []

            for saved_path in saved_files:
                parts = split_audio(saved_path, os.path.join(app.config["UPLOAD_FOLDER"], f"chunks_{job_id}_{os.path.splitext(os.path.basename(saved_path))[0]}"), chunk_length=60)
                total_chunks += len(parts)

            with JOBS_LOCK:
                JOBS[job_id]["status"] = "transcribing"
                JOBS[job_id]["total"] = total_chunks
                JOBS[job_id]["progress"] = 0

            transcriptor = Transcriptor()
            diarizador = Diarizador()
            formateador = Formateador()
            asignador_roles = AsignadorRoles()
            speaker_texts = {}
            audio_results_by_audio = {}

            for saved_path in saved_files:
                audio_name = os.path.splitext(os.path.basename(saved_path))[0]
                parts_dir = os.path.join(app.config["UPLOAD_FOLDER"], f"chunks_{job_id}_{audio_name}")
                parts = sorted([os.path.join(parts_dir, f) for f in os.listdir(parts_dir) if f.lower().endswith('.wav')])
                audio_results = []

                for part in parts:
                    r = transcriptor.transcribir(part)
                    diar = diarizador.diarizar(part)
                    assigned = formateador.unir(r, diar)
                    grouped = formateador.agrupar(assigned)

                    for block in grouped:
                        speaker = block.get("speaker", "SPEAKER_00")
                        if speaker not in speaker_set:
                            speaker_set.append(speaker)
                        speaker_texts.setdefault(speaker, []).append(block.get("texto", ""))
                        audio_results.append({
                            "audio": audio_name,
                            "speaker": speaker,
                            "text": block.get("texto", ""),
                        })

                    processed_chunks += 1
                    with JOBS_LOCK:
                        JOBS[job_id]["progress"] = processed_chunks

                all_results.extend(audio_results)
                audio_results_by_audio[audio_name] = audio_results

            transcriptor.liberar_memoria()

            speaker_texts_joined = {speaker: " ".join(texts) for speaker, texts in speaker_texts.items()}
            roles = asignador_roles.clasificar_hablantes(speaker_texts_joined)

            if "Entrevistador" not in roles.values() and speaker_set:
                roles[speaker_set[0]] = "Entrevistador"
            for speaker in speaker_set:
                roles.setdefault(speaker, "Entrevistado")

            entrevistado_ids = [s for s, r in roles.items() if r == "Entrevistado"]
            if len(entrevistado_ids) > 1:
                for idx, speaker in enumerate(entrevistado_ids, start=1):
                    roles[speaker] = f"Entrevistado {idx}"

            for item in all_results:
                item["role"] = roles.get(item["speaker"], item["speaker"])

            timestamp = time.strftime("%Y%m%d_%H%M%S")
            txt_name = f"Transcripcion_web_{timestamp}.txt"
            txt_path = os.path.join(OUTPUT_DIR, txt_name)
            with open(txt_path, "w", encoding="utf-8") as fh:
                for item in all_results:
                    fh.write(f"**{item['role']}**\n")
                    fh.write(item["text"] + "\n\n")

            for audio_name, audio_results in audio_results_by_audio.items():
                timestamp = time.strftime("%Y%m%d_%H%M%S")
                docx_name = f"{audio_name}_{timestamp}.docx"
                docx_path = os.path.join(OUTPUT_DIR, docx_name)
                doc = Document()
                doc.add_heading(audio_name, level=1)
                for item in audio_results:
                    item["role"] = roles.get(item["speaker"], item["speaker"])
                    p = doc.add_paragraph()
                    p.add_run(item["role"]).bold = True
                    p.add_run("\n")
                    p.add_run(item["text"])
                doc.save(docx_path)
                docx_names.append(docx_name)

            # Limpiar chunks después de terminar
            try:
                for saved_path in saved_files:
                    audio_name = os.path.splitext(os.path.basename(saved_path))[0]
                    chunks_dir = os.path.join(app.config["UPLOAD_FOLDER"], f"chunks_{job_id}_{audio_name}")
                    if os.path.exists(chunks_dir):
                        shutil.rmtree(chunks_dir)
                        print(f"Chunks eliminados: {chunks_dir}")
                    # También eliminar el archivo de audio subido
                    if os.path.exists(saved_path):
                        os.remove(saved_path)
                        print(f"Archivo de audio eliminado: {saved_path}")
            except Exception as cleanup_exc:
                print(f"Error al limpiar chunks: {cleanup_exc}")

            with JOBS_LOCK:
                JOBS[job_id]["status"] = "done"
                JOBS[job_id]["result"] = all_results
                JOBS[job_id]["txt_name"] = txt_name
                JOBS[job_id]["docx_names"] = docx_names
                JOBS[job_id]["roles"] = roles
        except Exception as exc:
            # Limpiar chunks incluso si hay error
            try:
                for saved_path in saved_files:
                    audio_name = os.path.splitext(os.path.basename(saved_path))[0]
                    chunks_dir = os.path.join(app.config["UPLOAD_FOLDER"], f"chunks_{job_id}_{audio_name}")
                    if os.path.exists(chunks_dir):
                        shutil.rmtree(chunks_dir)
                    if os.path.exists(saved_path):
                        os.remove(saved_path)
            except:
                pass
            
            with JOBS_LOCK:
                JOBS[job_id]["status"] = "error"
                JOBS[job_id]["error"] = str(exc)

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()

    # Devolver el job_id en lugar de redirigir
    return jsonify({"job_id": job_id, "status_url": f"/job-status/{job_id}"}), 202


@app.route('/job-status/<job_id>', methods=["GET", "OPTIONS"])
def job_status_json(job_id):
    if request.method == "OPTIONS":
        return "", 204
    
    with JOBS_LOCK:
        job = JOBS.get(job_id)
    if job is None:
        return jsonify({"error": "Job no encontrado"}), 404
    return jsonify(job)


@app.route('/salidas/<path:filename>', methods=["GET", "OPTIONS"])
def download_file(filename):
    if request.method == "OPTIONS":
        return "", 204
    
    if os.path.exists(os.path.join(OUTPUT_DIR, filename)):
        return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)
    return jsonify({"error": "Archivo no encontrado"}), 404


if __name__ == '__main__':
    import sys
    
    # Detectar si estamos en producción
    debug = '--debug' in sys.argv or os.environ.get('FLASK_DEBUG', '').lower() == 'true'
    
    # En producción, usar 0.0.0.0 para ser accesible desde fuera
    host = '127.0.0.1' if debug else '0.0.0.0'
    
    print(f"Iniciando TranscriptorIA API en {host}:5000")
    print(f"Modo debug: {debug}")
    app.run(host=host, port=5000, debug=debug, threaded=True)
